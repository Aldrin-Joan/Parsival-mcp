import os
import tempfile
import docx
import pytest
from src.tools.read_file import _read_file
from src.models.enums import OutputFormat, ParseStatus


def make_docx_file():
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    f.close()
    document = docx.Document()
    document.add_heading("Hello", level=1)
    document.add_paragraph("World")
    document.save(f.name)
    return f.name


@pytest.mark.asyncio
async def test_integration_read_file_happy_path():
    path = make_docx_file()
    try:
        result = await _read_file(path, output_format=OutputFormat.MARKDOWN)
        assert result.status == ParseStatus.OK
        assert "# Hello" in result.content
        assert "World" in result.content
        assert result.cache_hit is False or isinstance(result.cache_hit, bool)
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_integration_read_file_unsupported_format(tmp_path):
    path = tmp_path / "binary.bin"
    path.write_bytes(b"\x00\x01\x02\x03")

    result = await _read_file(str(path), output_format=OutputFormat.MARKDOWN)
    assert result.status == ParseStatus.UNSUPPORTED
    assert result.content == ""
    assert len(result.errors) == 1
    assert result.errors[0].code == "unsupported_format"


@pytest.mark.asyncio
async def test_integration_markdown_output_contract(tmp_path):
    path = tmp_path / "text.md"
    path.write_text("# Title\n\nA paragraph.")

    result = await _read_file(str(path), output_format=OutputFormat.MARKDOWN)
    assert result.status == ParseStatus.OK
    assert result.content.startswith("---")
    assert "source:" in result.content
    assert "format:" in result.content
    assert "generated_at" in result.content
    assert "# Title" in result.content


@pytest.mark.asyncio
async def test_integration_get_metadata_flow(tmp_path):
    from src.tools.get_metadata import get_metadata

    path = tmp_path / "plain.txt"
    path.write_text("Hello world")

    metadata = await get_metadata(str(path))
    assert metadata.file_format == "text"
    assert metadata.source_path == str(path)
    assert metadata.word_count == 2
    assert metadata.section_count >= 0
