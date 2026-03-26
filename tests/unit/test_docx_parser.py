import os
import tempfile
from pathlib import Path
import docx
import pytest
from src.parsers.docx_parser import DocxParser
from src.models.enums import ParseStatus, SectionType


def make_docx_file():
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    doc = docx.Document()
    doc.add_heading('Title', level=1)
    p = doc.add_paragraph('Hello world')
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'A'
    table.rows[0].cells[1].text = 'B'
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = '2'
    doc.save(tmp.name)
    return tmp.name


def test_docx_parse_order_and_content():
    path = make_docx_file()
    try:
        parser = DocxParser()
        result = __import__('asyncio').run(parser.parse(Path(path)))

        assert result.status == ParseStatus.OK
        assert len(result.sections) >= 2
        assert result.sections[0].type == SectionType.HEADING
        assert 'Title' in result.sections[0].content
        assert any(s.type == SectionType.PARAGRAPH for s in result.sections)
        assert result.metadata.table_count == 1
        assert result.metadata.image_count == 0

        table = result.tables[0]
        assert table.headers == ['A', 'B']
        assert table.row_count == 1
        assert table.col_count == 2
    finally:
        os.unlink(path)


def test_docx_parse_metadata():
    path = make_docx_file()
    try:
        parser = DocxParser()
        metadata = __import__('asyncio').run(parser.parse_metadata(Path(path)))
        assert metadata.source_path == path
        assert metadata.file_format == 'docx'
        assert metadata.page_count is None
    finally:
        os.unlink(path)



@pytest.mark.asyncio
async def test_docx_parser_encrypted_file():
    from zipfile import ZipFile

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()

    with ZipFile(tmp.name, 'w') as zf:
        zf.writestr('[Content_Types].xml', '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>')
        zf.writestr('EncryptionInfo', 'some dummy encryption info')
        zf.writestr('EncryptedPackage', 'encrypted-data')

    try:
        parser = DocxParser()
        result = await parser.parse(Path(tmp.name))

        assert result.status == ParseStatus.FAILED
        assert result.errors
        assert result.errors[0].code == 'encrypted'
    finally:
        os.unlink(tmp.name)
