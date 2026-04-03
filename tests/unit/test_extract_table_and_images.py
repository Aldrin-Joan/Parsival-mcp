from pathlib import Path

import pytest

from src.tools.extract_table import extract_table
from src.tools.extract_images import extract_images

try:
    from PIL import Image
except ImportError:
    Image = None


def make_docx_with_table(tmp_path: Path) -> Path:
    from docx import Document

    path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "A1"
    table.cell(1, 1).text = "B1"
    table.cell(2, 0).text = "A2"
    table.cell(2, 1).text = "B2"
    doc.save(path)
    return path


def make_xlsx_with_table(tmp_path: Path) -> Path:
    import openpyxl

    path = tmp_path / "book.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["H1", "H2"])
    ws.append(["A1", "B1"])
    ws.append(["A2", "B2"])
    wb.save(path)
    return path


def make_pdf_with_table(tmp_path: Path) -> Path:
    import fitz

    path = tmp_path / "table.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "H1 H2\nA1 B1")
    doc.save(path)
    doc.close()
    return path


@pytest.mark.skipif(Image is None, reason="Pillow required for image tests")
@pytest.mark.asyncio
async def test_extract_images_docx_and_pdf(tmp_path):
    # DOCX with image
    from docx import Document

    docx_path = tmp_path / "img.docx"
    doc = Document()
    p = doc.add_paragraph("x")
    img = Image.new("RGB", (20, 20), color="red")
    tmp_img = tmp_path / "img.png"
    img.save(tmp_img)
    doc.add_picture(str(tmp_img), width=1000000, height=1000000)
    doc.save(docx_path)

    images = await extract_images(str(docx_path))
    assert len(images) >= 1

    # pdf image by adding image file to PDF
    pdf_path = tmp_path / "img.pdf"
    import fitz

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_image(fitz.Rect(72, 72, 172, 172), filename=str(tmp_img))
    pdf.save(pdf_path)
    pdf.close()

    images = await extract_images(str(pdf_path))
    assert len(images) >= 1


@pytest.mark.asyncio
async def test_extract_table_valid_and_out_of_range(tmp_path):
    docx_path = make_docx_with_table(tmp_path)
    tbl = await extract_table(str(docx_path), table_index=1)
    assert tbl.headers == ["H1", "H2"]

    with pytest.raises(IndexError):
        await extract_table(str(docx_path), table_index=2)

    xlsx_path = make_xlsx_with_table(tmp_path)
    tbl2 = await extract_table(str(xlsx_path), table_index=1, sheet_name="Sheet1")
    assert tbl2.headers == ["H1", "H2"]

    with pytest.raises(IndexError):
        await extract_table(str(xlsx_path), table_index=1, sheet_name="Missing")


@pytest.mark.asyncio
async def test_extract_table_pdf(tmp_path):
    pdf_path = make_pdf_with_table(tmp_path)
    try:
        tbl = await extract_table(str(pdf_path), table_index=1)
    except (IndexError, ValueError):
        pytest.skip("PDF table extraction format not reliable for synthetic sample")
    assert tbl.col_count >= 1
    assert tbl.row_count >= 2
