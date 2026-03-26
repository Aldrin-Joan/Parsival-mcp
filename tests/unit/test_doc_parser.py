import asyncio
import os
import sys
from pathlib import Path
import shutil
import tempfile

import pytest

try:
    import docx
except ImportError:
    docx = None

from src.models.enums import FileFormat, ParseStatus
from src.parsers.doc_parser import DocParser

try:
    import psutil
except ImportError:
    psutil = None


def create_docx_file(path: Path):
    from docx import Document

    doc = Document()
    doc.add_heading('Hello From DOC', level=1)
    doc.add_paragraph('This is a test DOCX-based content.')
    doc.save(path)


@pytest.mark.skipif(docx is None, reason='python-docx is required for these tests')
@pytest.mark.asyncio
async def test_doc_parser_valid_doc(monkeypatch, tmp_path):
    base = tmp_path
    source_doc = base / 'sample.doc'
    converted_docx = base / 'sample.docx'

    create_docx_file(converted_docx)
    shutil.copy(converted_docx, source_doc)

    async def fake_convert(self, path):
        assert path == source_doc
        temp_dir = base / 'tmpconv'
        temp_dir.mkdir(exist_ok=True)
        return converted_docx, temp_dir

    monkeypatch.setattr(DocParser, '_convert_doc_to_docx', fake_convert)

    parser = DocParser()
    result = await parser.parse(source_doc)

    assert result.status == ParseStatus.OK
    assert result.metadata.file_format == FileFormat.DOC
    assert result.metadata.source_path == str(source_doc)
    assert len(result.sections) >= 1
    assert 'Hello From DOC' in result.raw_text


@pytest.mark.skipif(docx is None, reason='python-docx is required for these tests')
@pytest.mark.asyncio
async def test_doc_parser_corrupt_doc(monkeypatch, tmp_path):
    source_doc = tmp_path / 'corrupt.doc'
    source_doc.write_text('not-a-valid-doc', encoding='utf-8')

    async def fake_convert(self, path):
        raise RuntimeError('conversion failed: corrupt document')

    monkeypatch.setattr(DocParser, '_convert_doc_to_docx', fake_convert)

    parser = DocParser()
    result = await parser.parse(source_doc)

    assert result.status == ParseStatus.FAILED
    assert result.errors and result.errors[0].code == 'conversion_failed'


@pytest.mark.skipif(docx is None, reason='python-docx is required for these tests')
@pytest.mark.asyncio
async def test_doc_parser_timeout(monkeypatch, tmp_path):
    source_doc = tmp_path / 'timeout.doc'
    source_doc.write_text('garbage', encoding='utf-8')

    async def fake_convert(self, path):
        raise TimeoutError('LibreOffice conversion timed out')

    monkeypatch.setattr(DocParser, '_convert_doc_to_docx', fake_convert)

    parser = DocParser()
    result = await parser.parse(source_doc)

    assert result.status == ParseStatus.FAILED
    assert result.errors and result.errors[0].code == 'conversion_timeout'


@pytest.mark.skipif(docx is None, reason='python-docx is required for these tests')
@pytest.mark.asyncio
async def test_doc_parser_concurrent_semaphore(monkeypatch, tmp_path):
    source_doc = tmp_path / 'concurrent.doc'
    converted_docx = tmp_path / 'concurrent.docx'
    create_docx_file(converted_docx)
    source_doc.write_text('dummy', encoding='utf-8')

    active = 0
    maximum = 0

    async def fake_convert(self, path):
        nonlocal active, maximum
        active += 1
        maximum = max(maximum, active)
        await asyncio.sleep(0.1)
        active -= 1
        temp_dir = tmp_path / 'tmpconv' / str(os.getpid())
        temp_dir.mkdir(parents=True, exist_ok=True)
        return converted_docx, temp_dir

    monkeypatch.setattr(DocParser, '_convert_doc_to_docx', fake_convert)

    parser = DocParser()
    tasks = [parser.parse(source_doc) for _ in range(6)]
    results = await asyncio.gather(*tasks)

    assert all(r.status == ParseStatus.OK for r in results)
    assert maximum <= 2


@pytest.mark.skipif(psutil is None, reason='psutil is required for this test')
@pytest.mark.asyncio
async def test_doc_parser_subprocess_timeout_kills_orphans(tmp_path):
    parser = DocParser()
    cmd = [
        sys.executable,
        '-c',
        'import time; time.sleep(10)',
    ]

    before = {p.pid for p in psutil.process_iter(attrs=['pid', 'cmdline']) if p.info['cmdline'] and sys.executable in p.info['cmdline']}

    from src.parsers import doc_parser

    with pytest.raises(TimeoutError):
        await doc_parser._run_subprocess(cmd, timeout=0.2)

    await asyncio.sleep(0.2)
    after = {p.pid for p in psutil.process_iter(attrs=['pid', 'cmdline']) if p.info['cmdline'] and sys.executable in p.info['cmdline']}

    assert after.issubset(before)


@pytest.mark.asyncio
async def test_executor_parse_in_pool_concurrency(tmp_path):
    from src.core.executor import run_parse_in_pool
    from src.models.enums import OutputFormat

    p = tmp_path / 'p.txt'
    p.write_text('hello world\n' * 1000, encoding='utf-8')

    tasks = [run_parse_in_pool(str(p), options={'output_format': OutputFormat.JSON.value}) for _ in range(6)]
    results = await asyncio.gather(*tasks)

    assert all(r.status == ParseStatus.OK for r in results)

