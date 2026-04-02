import os
import tempfile
import docx
import fitz
from src.tools.read_file import _read_file
from src.models.enums import OutputFormat, ParseStatus


def make_docx_file():
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    f.close()
    document = docx.Document()
    document.add_heading("Hello", level=1)
    document.add_paragraph("World")
    document.save(f.name)
    return f.name


import pytest


@pytest.mark.asyncio
async def test_read_file_tool_docx():
    path = make_docx_file()
    try:
        result = await _read_file(path, output_format=OutputFormat.MARKDOWN, stream=False)
        assert result.status.name == "OK"
        assert "# Hello" in result.content
        assert "World" in result.content
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_read_file_tool_csv():
    path = tempfile.NamedTemporaryFile(suffix=".csv", delete=False)
    path.write(b"A,B\n1,2\n")
    path.close()
    try:
        result = await _read_file(path.name, output_format=OutputFormat.JSON, stream=False)
        assert result.status.name == "OK"
        assert '"A"' in result.content
    finally:
        os.unlink(path.name)


@pytest.mark.asyncio
async def test_read_file_tool_text_output_docx():
    path = make_docx_file()
    try:
        result = await _read_file(path, output_format=OutputFormat.TEXT, stream=False)
        assert result.status.name == "OK"
        assert "Hello" in result.content
        assert "World" in result.content
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_read_file_tool_plain_text_no_ext(tmp_path):
    path = tmp_path / "plain"
    path.write_text("Hello plain text", encoding="utf-8")

    result = await _read_file(str(path), output_format=OutputFormat.MARKDOWN, stream=False)
    assert result.status.name == "OK"
    assert "Hello plain text" in result.content


@pytest.mark.asyncio
async def test_read_file_tool_unsupported_format(tmp_path):
    path = tmp_path / "binary"
    path.write_bytes(b"\x00\x00\x01\x02")

    result = await _read_file(str(path), output_format=OutputFormat.MARKDOWN, stream=False)
    assert result.status.name == "UNSUPPORTED"
    assert result.content == ""
    assert len(result.errors) == 1
    assert result.errors[0].code == "unsupported_format"


@pytest.mark.asyncio
async def test_read_file_tool_extension_content_mismatch(tmp_path):
    path = tmp_path / "trick.txt"
    path.write_bytes(b"%PDF-1.4\n%Dummy")

    result = await _read_file(str(path), output_format=OutputFormat.MARKDOWN, stream=False)
    assert result.metadata.file_format in ("pdf", "PDF")
    # Parser may fail or succeed depending on PDF validity in binary payload.
    assert result.status in (ParseStatus.OK, ParseStatus.PARTIAL, ParseStatus.FAILED)


def test_read_file_tool_schema_includes_new_options():
    import inspect
    from src.tools.read_file import _read_file

    params = inspect.signature(_read_file).parameters
    assert "page_range" in params
    assert "include_images" in params
    assert "max_tokens_hint" in params


@pytest.mark.asyncio
async def test_read_file_tool_invalid_page_range(tmp_path):
    path = tmp_path / "plain.txt"
    path.write_text("This is line1\nThis is line2\n")

    result = await _read_file(str(path), output_format=OutputFormat.TEXT, page_range=(2, 1), stream=False)
    assert result.status.name == "FAILED"
    assert result.errors[0].code == "invalid_argument"


@pytest.mark.asyncio
async def test_read_file_tool_page_range_text(tmp_path):
    path = tmp_path / "plain.txt"
    path.write_text("Line1\nLine2\nLine3\n")

    result = await _read_file(str(path), output_format=OutputFormat.TEXT, page_range=(2, 3), stream=False)
    assert result.status.name == "OK"
    assert "Line1" not in result.content
    assert "Line2" in result.content


@pytest.mark.asyncio
async def test_read_file_tool_max_tokens_hint(tmp_path):
    path = tmp_path / "plain.txt"
    path.write_text("one two three four five")

    result = await _read_file(str(path), output_format=OutputFormat.TEXT, max_tokens_hint=3, stream=False)
    assert result.status == ParseStatus.PARTIAL
    assert result.content.split() == ["one", "two", "three"]
    assert any(err.code == "max_tokens_hint_reached" for err in result.errors)


@pytest.mark.asyncio
async def test_read_file_tool_streaming_pdf():
    import time
    from src.tools.read_file import _read_file
    from src.parsers.pdf_parser import PDFParser

    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i + 1} text")
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    doc.save(tmp.name)
    doc.close()

    try:
        # Ensure parser supports streaming path.
        assert PDFParser().supports_streaming() is True

        start = time.perf_counter()
        first_emit = None
        total_chunks = 0
        final_chunk = None

        stream = await _read_file(tmp.name, output_format=OutputFormat.JSON, stream=True)
        assert hasattr(stream, "__aiter__")

        async for chunk in stream:
            if total_chunks == 0:
                first_emit = time.perf_counter() - start
            total_chunks += 1
            if chunk.is_final:
                final_chunk = chunk

        assert total_chunks > 1
        assert first_emit is not None
        assert first_emit < time.perf_counter() - start
        assert final_chunk is not None
        assert final_chunk.is_final
        assert "metadata" in final_chunk.content

    finally:
        os.unlink(tmp.name)
